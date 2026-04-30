import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not found. Please run: pip3 install python-docx")
    exit(1)

def set_font(run, font_name_cn='宋体', font_name_en='Times New Roman', font_size=10.5):
    """Set Chinese and English fonts for a run."""
    run.font.size = Pt(font_size)
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)

def create_environmental_paper(title, authors, abstract, keywords, body_chunks, references, output_path):
    doc = Document()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_font(run, font_name_cn='黑体', font_size=18) # 2nd size
    
    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors)
    set_font(run, font_size=12)
    
    # Abstract
    p = doc.add_paragraph()
    run = p.add_run("摘要：")
    set_font(run, font_name_cn='黑体', font_size=10.5)
    run = p.add_run(abstract)
    set_font(run, font_size=10.5)
    
    # Keywords
    p = doc.add_paragraph()
    run = p.add_run("关键词：")
    set_font(run, font_name_cn='黑体', font_size=10.5)
    run = p.add_run(keywords)
    set_font(run, font_size=10.5)
    
    # Body
    for heading, content in body_chunks:
        # Heading
        h = doc.add_paragraph()
        run = h.add_run(heading)
        set_font(run, font_name_cn='黑体', font_size=14)
        
        # Content
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(21) # 2 characters
                run = p.add_run(line.strip())
                set_font(run, font_size=10.5)
    
    # References
    doc.add_page_break()
    p = doc.add_heading("参考文献", level=1)
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        set_font(run, font_size=9)
        
    doc.save(output_path)
    print(f"Paper generated successfully at: {output_path}")

if __name__ == "__main__":
    # Sample execution
    create_environmental_paper(
        title="环境领域AIGC辅助研究的现状与挑战",
        authors="陈无恙",
        abstract="本文探讨了人工智能生成内容（AIGC）在环境科学研究中的应用...",
        keywords="环境科学；AIGC；查重；知网",
        body_chunks=[
            ("1. 引言", "随着大数据技术的发展，环境领域的研究日益复杂..."),
            ("2. 研究方法", "基于知网（CNKI）与维普等平台的文献调研..."),
        ],
        references=[
            "生态环境部. 十四五生态环境保护规划[R]. 2021.",
            "张三. 环境科学研究方法[M]. 北京: 科学出版社, 2023."
        ],
        output_path="environmental_sample_draft.docx"
    )
