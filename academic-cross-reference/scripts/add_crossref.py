#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学术文档交叉引用处理器
为 Word 文档添加参考文献角标与参考文献列表之间的双向超链接

用法:
    python add_crossref.py <输入文件.docx> <输出文件.docx>
"""

import sys
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re

def add_cross_references(input_path, output_path):
    """
    为 Word 文档添加交叉引用：
    1. 识别参考文献列表中的条目
    2. 为每个参考文献添加书签
    3. 将正文中的角标 [1] [2] 等转换为可点击的超链接，指向对应参考文献
    """
    doc = Document(input_path)
    
    # 步骤 1: 找到参考文献列表
    ref_indices = []
    in_refs = False
    ref_start = None
    
    for i, p in enumerate(doc.paragraphs):
        if not in_refs and '参考文献' in p.text and len(p.text.strip()) <= 6:
            in_refs = True
            ref_start = i
            continue
        if in_refs and p.text.strip():
            m = re.match(r'^\[(\d+)\]', p.text.strip())
            if m:
                ref_indices.append((i, int(m.group(1))))
    
    if not ref_indices:
        print("警告: 未找到参考文献列表")
        return
    
    print(f"找到 {len(ref_indices)} 篇参考文献，起始段落: {ref_start}")
    print(f"参考文献编号: {[r[1] for r in ref_indices]}")
    
    # 步骤 2: 为每个参考文献添加书签
    for para_idx, ref_num in ref_indices:
        p = doc.paragraphs[para_idx]
        # 添加书签开始
        bs = parse_xml('<w:bookmarkStart {} w:id="{}" w:name="ref{}"/>'.format(
            nsdecls('w'), ref_num, ref_num))
        # 添加书签结束
        be = parse_xml('<w:bookmarkEnd {} w:id="{}"/>'.format(
            nsdecls('w'), ref_num))
        p._p.insert(0, bs)
        p._p.append(be)
    
    # 步骤 3: 将正文中的角标转换为超链接
    count = 0
    for i, p in enumerate(doc.paragraphs):
        # 跳过参考文献列表本身
        if ref_start is not None and i >= ref_start:
            break
        
        text = p.text
        for m in re.finditer(r'\[(\d+)\]', text):
            ref_num = int(m.group(1))
            # 只处理存在于参考文献列表中的编号
            if any(r[1] == ref_num for r in ref_indices):
                part = p.part
                # 创建超链接关系
                part.relate_to("ref{}".format(ref_num), RT.HYPERLINK, is_external=True)
                
                # 创建超链接 XML
                ns = nsdecls('w')
                hx = '<w:hyperlink %s w:anchor="ref%d" w:history="1">' \
                     '<w:r %s>' \
                     '<w:rPr>' \
                     '<w:rFonts %s w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体"/>' \
                     '<w:sz w:val="16"/>' \
                     '<w:szCs w:val="16"/>' \
                     '<w:vertAlign w:val="superscript"/>' \
                     '<w:color w:val="0563C1"/>' \
                     '</w:rPr>' \
                     '<w:t xml:space="preserve">[%d]</w:t>' \
                     '</w:r>' \
                     '</w:hyperlink>' % (ns, ref_num, ns, ns, ref_num)
                
                p._p.append(parse_xml(hx))
                count += 1
    
    print(f"成功添加 {count} 个交叉引用超链接")
    doc.save(output_path)
    print(f"完成! 输出文件: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python add_crossref.py <输入文件.docx> <输出文件.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    add_cross_references(input_file, output_file)
